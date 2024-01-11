from bs4 import BeautifulSoup
from urllib.parse import urljoin
from tabulate import tabulate
from selenium import webdriver

from docx import Document
from docx.shared import Inches

import io
import os
import requests
from PIL import Image
import ast

from .constants import(
    SPECIAL_CONTAINER_TAGS,
    SPLIT_LINES,
    FILE_EXTENSION,
    PARAGRAPH_TAGS,
    STATIC_FILES_OUTPUT_DIRECTORY,
    OUTPUT_FILES_DIRECTORY,
    HEADER_TAGS,
    # HTML_CONTENT
    
)

doc = Document()
doc.add_heading('Centrix Web-Scraping Application', level=1)
chrome_driver = webdriver.Chrome()
class SearchWithContainers:
    def __init__(self, domain_name, container_tag='body', output_filename='web_scrape'):
        self.domain_name = domain_name
        self.container_tag = container_tag
        self.soup_object = self.get_soup_object()
        self.output_file = output_filename + FILE_EXTENSION

    def get_soup_object(self):
        html_content = requests.get(self.domain_name).text
        html_content = self.get_webpage_content()
        # html_content = HTML_CONTENT
        soup = BeautifulSoup(html_content, 'lxml')
        return soup
    
    def get_webpage_content(self):
        chrome_driver.execute_script("window.open('','_blank');")
        chrome_driver.switch_to.window(chrome_driver.window_handles[-1])
        chrome_driver.get(self.domain_name)

        """
        html_content = chrome_driver.page_source
        input_filename = os.path.join(INPUT_FILES_DIRECTORY, 'web_content.html')
        with open(input_filename, 'w', encoding='utf-8') as file:
            file.write(html_content)
        """
        html_content = chrome_driver.page_source
        # Close the browser window
        # chrome_driver.quit()
        return html_content

    def getTagsBKP(self):
        tags_dict = {}
        tags_data = [tag for tag in self.container_tag_data.find_all()]
        count = 1
        for tag in tags_data:
            tag_name = tag.name
            if tag_name not in tags_dict:
                tags_dict[tag_name] = []
            tags_dict[tag_name].append(tag)
            count += 1
            if count == 5:
                break
        return tags_dict
    
    def get_div_content(self, contents):
        text_content = ''
        for content in contents:
            try:
                text = content.strip()
                if text.startswith('<') and text.endswith('/>'):
                    break
                else:
                    text_content += text
            except Exception as e:
                break
        return text_content
                
    def get_all_tags(self):
        tags_dict = {}
        tags_data = self.soup_object.find_all()
        paragraph_flag = True

        for tag in tags_data:
            tag_name = tag.name
            tag_text = tag.get_text(strip=True)
            if tag_name in ['div']:
                contents = tag.contents
                tag_text = self.get_div_content(contents)

            if tag_name not in tags_dict:
                if tag_name in PARAGRAPH_TAGS and paragraph_flag:
                    tags_dict['p'] = []
                    paragraph_flag = False
                else:
                    tags_dict[tag_name] = []
                    
            if tag_name == 'img':
                tag_text = tag.get('src', '')
                tags_dict[tag_name].append({tag_name: tag_text})

            elif tag_name in PARAGRAPH_TAGS and tag_text:
                tags_dict['p'].append({tag_name: tag_text})

        return tags_dict

    def run_search_tags(self):
        return self.get_all_tags()    

class CreateWordDocument:
    def __init__(self, tags_data):
        self.tags_data = tags_data

    #NOTE:: To handle huge requests we need to write this logic in async after sometime
    def process_image_tags(self, img_url):
        # img_url = urljoin(self.domain_name, img_url)
        img_data = requests.get(img_url).content
        img_filename = os.path.join(STATIC_FILES_OUTPUT_DIRECTORY, os.path.basename(img_url))

        #Save images in local directory
        with open(img_filename, 'wb') as img_file:
            img_file.write(img_data)
            # print(f"Image saved in local directory")

        #Open and write to docx in png format
        with Image.open(img_filename) as img_file:
            img_byte_array = io.BytesIO()
            img_file.save(img_byte_array, format='PNG')
            doc.add_picture(img_byte_array, width=Inches(2))


    def insert_into_word_document(self, tag_data):
        tag_key = ''
        tag_content = ''
        # {'h1': 'Madhava Krishna Educational Society'}
        for tag, content in tag_data.items():
            tag_key = tag
            tag_content = content
            break

        if tag_key == 'p':
            doc.add_paragraph(tag_content)
        elif tag_key in HEADER_TAGS:
            if tag_key == 'h1':
                doc.add_heading(tag_content, level=1)
            elif tag_key == 'h2':
                doc.add_heading(tag_content, level=2)
            elif tag_key == 'h3':
                doc.add_heading(tag_content, level=3)
            elif tag_key == 'h4':
                doc.add_heading(tag_content, level=4)
            elif tag_key == 'h5':
                doc.add_heading(tag_content, level=5)
            elif tag_key == 'h6':
                doc.add_heading(tag_content, level=6)
        
        # elif tag_key == 'img':
            # self.process_image_tags(tag_content)

    def run_create_word_document(self):
        # selected_tags = ["{'h1': 'Madhava Krishna Educational Society'}"]
        for tag_data in self.tags_data:
            #convert string to dict
            tag_dict = ast.literal_eval(tag_data)
            self.insert_into_word_document(tag_dict)
        
        filename = "report-webscrap.docx"
        output_files_path = os.path.join(OUTPUT_FILES_DIRECTORY, filename)
        print(f"output_files_path: {output_files_path}")
        doc.save(output_files_path)


# website_url = 'http://makes.org.in/'
# container_tag = 'body'
# output_filename = 'test'

# scrap_object = SearchWithContainers(website_url, container_tag, output_filename)
# result = scrap_object.run_search_tags()
# print(f'result:', result)
